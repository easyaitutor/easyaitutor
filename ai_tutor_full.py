from pathlib import Path
from dotenv import load_dotenv
import os, io, json, traceback, re, uuid, random, mimetypes, string, csv
from datetime import date, datetime, timedelta, timezone as dt_timezone
import openai
import gradio as gr
from docx import Document
import smtplib
from email.message import EmailMessage
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import RedirectResponse, HTMLResponse
import jwt
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
from fastapi.middleware.cors import CORSMiddleware
from fastapi.templating import Jinja2Templates


# --- Load environment variables ---
dotenv_path = Path(__file__).parent / ".env"
load_dotenv(dotenv_path=dotenv_path)

# --- Configuration ---
openai.api_key = os.getenv("OPENAI_API_KEY")
CONFIG_DIR = Path("course_data")
CONFIG_DIR.mkdir(exist_ok=True)
PROGRESS_LOG_FILE = CONFIG_DIR / "student_progress_log.csv"

SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT   = int(os.getenv("SMTP_PORT", 587))
SMTP_USER   = os.getenv("SMTP_USER")
SMTP_PASS   = os.getenv("SMTP_PASS")
SUPPORT_EMAIL_ADDRESS = os.getenv("SUPPORT_EMAIL_ADDRESS", "easyaitutor@gmail.com")

JWT_SECRET_KEY = os.getenv("JWT_SECRET_KEY", "a-very-secure-secret-key-please-change")
if JWT_SECRET_KEY == "a-very-secure-secret-key-please-change":
    print("WARNING: JWT_SECRET_KEY is set to default. Set a strong secret key in env variables.")
LINK_VALIDITY_HOURS = int(os.getenv("LINK_VALIDITY_HOURS", 24))
ALGORITHM = "HS256"
APP_DOMAIN = os.getenv("APP_DOMAIN", "https://www.easyaitutor.com") # Ensure this matches audience in JWT

days_map = {"Monday":0, "Tuesday":1, "Wednesday":2, "Thursday":3,
            "Friday":4, "Saturday":5, "Sunday":6}

# --- Student Tutor Configuration ---
STUDENT_TTS_MODEL = "tts-1"
STUDENT_CHAT_MODEL = "gpt-4o-mini"
STUDENT_WHISPER_MODEL = "whisper-1"
STUDENT_DEFAULT_ENGLISH_LEVEL = "B1 (Intermediate)"
STUDENT_AUDIO_DIR = Path("student_audio_files")
STUDENT_AUDIO_DIR.mkdir(exist_ok=True)
STUDENT_BOT_NAME = "Easy AI Tutor"
STUDENT_LOGO_PATH = "logo.png" # Ensure this path is accessible or remove
STUDENT_ONBOARDING_TURNS = 2
STUDENT_TEACHING_TURNS_PER_BREAK = 5
STUDENT_INTEREST_BREAK_TURNS = 1
STUDENT_QUIZ_AFTER_TURNS = 7
STUDENT_MAX_SESSION_TURNS = 20
STUDENT_UI_PATH = "/student_tutor_interface/ui" # Used for mounting and redirect

EASYAI_TUTOR_PROGRESS_API_ENDPOINT = os.getenv("EASYAI_TUTOR_PROGRESS_API_ENDPOINT") # Less used now


# ─── Create FastAPI app & CORS ───────────────────────────────────────────────
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

# --- APScheduler Setup & Jobs ---
scheduler = BackgroundScheduler(timezone="UTC")

# Health‐check endpoint
@app.get("/healthz")
def healthz():
    return {"status": "ok", "scheduler_running": scheduler.running}

# Attempt to import fitz (PyMuPDF)
try:
    import fitz
    fitz_available = True
except ImportError:
    fitz_available = False
    print("PyMuPDF (fitz) not found. Page number mapping will be limited.")

# --- PDF Processing & Helpers ---
def split_sections(pdf_file_obj):
    if hasattr(pdf_file_obj, "seek"): pdf_file_obj.seek(0)
    if fitz_available:
        try:
            doc = None
            if hasattr(pdf_file_obj, "name"): doc = fitz.open(pdf_file_obj.name)
            elif hasattr(pdf_file_obj, "read"):
                pdf_bytes_sec = pdf_file_obj.read(); pdf_file_obj.seek(0)
                doc = fitz.open(stream=pdf_bytes_sec, filetype="pdf")
            if not doc: raise Exception("Could not open PDF with fitz.")
            pages_text = [page.get_text("text", sort=True) for page in doc]; doc.close()
            headings = []
            for i, text in enumerate(pages_text):
                for m in re.finditer(r"(?im)^(?:CHAPTER|Cap[ií]tulo|Sección|Section|Unit|Unidad|Part|Module)\s+[\d\w]+.*", text):
                    headings.append({"page": i + 1, "start_char_index": m.start(), "title": m.group().strip(), "page_index": i})
            headings.sort(key=lambda h: (h['page_index'], h['start_char_index']))
            sections = []
            if not headings:
                full_content = "\n".join(pages_text)
                if full_content.strip(): sections.append({'title': 'Full Document Content', 'content': full_content.strip(), 'page': 1})
                return sections
            for idx, h in enumerate(headings):
                current_page_idx, start_char = h['page_index'], h['start_char_index']; content = ''
                if idx + 1 < len(headings):
                    next_h = headings[idx+1]; next_page_idx, end_char = next_h['page_index'], next_h['start_char_index']
                    if current_page_idx == next_page_idx: content += pages_text[current_page_idx][start_char:end_char]
                    else:
                        content += pages_text[current_page_idx][start_char:] + '\n'
                        for p_idx in range(current_page_idx + 1, next_page_idx): content += pages_text[p_idx] + '\n'
                        content += pages_text[next_page_idx][:end_char]
                else:
                    content += pages_text[current_page_idx][start_char:] + '\n'
                    for p_idx in range(current_page_idx + 1, len(pages_text)): content += pages_text[p_idx] + '\n'
                if content.strip(): sections.append({'title': h['title'], 'content': content.strip(), 'page': h['page']})
            sections = [s for s in sections if len(s['content']) > len(s['title']) + 20]
            return sections
        except Exception as e_fitz: print(f"Error fitz splitting: {e_fitz}. Fallback.");
    try:
        from PyPDF2 import PdfReader
        if hasattr(pdf_file_obj, "seek"): pdf_file_obj.seek(0)
        reader = PdfReader(pdf_file_obj.name if hasattr(pdf_file_obj, "name") else pdf_file_obj)
        text = "\n".join(page.extract_text() or '' for page in reader.pages)
        chunks, sections, sents_per_sec = re.split(r'(?<=[.?!])\s+', text), [], 15
        for i in range(0, len(chunks), sents_per_sec):
            title, content = f"Content Block {i//sents_per_sec+1}", " ".join(chunks[i:i+sents_per_sec]).strip()
            if content: sections.append({'title': title, 'content': content, 'page': None})
        if not sections and text.strip(): sections.append({'title': 'Full Document (PyPDF2)', 'content': text.strip(), 'page': None})
        return sections
    except ImportError: return [{'title': 'PDF Error', 'content': 'PyPDF2 not found.', 'page': None}]
    except Exception as e_pypdf2: return [{'title': 'PDF Error', 'content': f'{e_pypdf2}', 'page': None}]

def download_docx(content, filename):
    buf = io.BytesIO(); doc = Document()
    for line in content.split("\n"):
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'): p.add_run(part[2:-2]).bold = True
            else: p.add_run(part)
    doc.save(buf); buf.seek(0); return buf, filename

def count_classes(sd, ed, wdays):
    cnt, cur = 0, sd
    while cur <= ed:
        if cur.weekday() in wdays: cnt += 1
        cur += timedelta(days=1)
    return cnt

def generate_access_token(student_id, course_id, lesson_id, lesson_date_obj=None):
    access_code = generate_5_digit_code()
    now = datetime.now(dt_timezone.utc)
    exp = now + timedelta(hours=LINK_VALIDITY_HOURS)
    payload = {
        "sub": student_id,
        "course_id": course_id,
        "lesson_id": lesson_id,
        "code": access_code,
        "iat": now,
        "exp": exp,
        "aud": APP_DOMAIN
    }
    token = jwt.encode(payload, JWT_SECRET_KEY, algorithm=ALGORITHM)
    return token, access_code

def generate_5_digit_code(): return str(random.randint(10000, 99999))

def send_email_notification(to_email, subject, html_content, from_name="User", attachment_file_obj=None):
    if not SMTP_USER or not SMTP_PASS: print(f"CRITICAL SMTP ERROR: SMTP_USER or SMTP_PASS not configured. Cannot send email to {to_email}."); return False
    msg = EmailMessage(); msg["Subject"] = subject; msg["From"] = f"AI Tutor Panel <{SMTP_USER}>"; msg["To"] = to_email
    if to_email.lower() == SMTP_USER.lower() and "@" in from_name: msg.add_header('Reply-To', from_name)
    msg.add_alternative(html_content, subtype='html')
    if attachment_file_obj and hasattr(attachment_file_obj, "name") and attachment_file_obj.name:
        try:
            # If attachment_file_obj is a Gradio File object, its .name is the path
            file_path_to_read = attachment_file_obj.name
            with open(file_path_to_read, 'rb') as fp: file_data = fp.read()
            ctype, encoding = mimetypes.guess_type(file_path_to_read)
            if ctype is None or encoding is not None: ctype = 'application/octet-stream'
            maintype, subtype_val = ctype.split('/', 1)
            msg.add_attachment(file_data,maintype=maintype,subtype=subtype_val,filename=os.path.basename(file_path_to_read))
            print(f"Attachment {os.path.basename(file_path_to_read)} prepared.")
        except FileNotFoundError: print(f"Error attaching: File not found at {file_path_to_read}")
        except Exception as e_attach: print(f"Error processing attachment {file_path_to_read}: {e_attach}")
    try:
        print(f"Attempting to send email to {to_email} via {SMTP_SERVER}:{SMTP_PORT} as {SMTP_USER}...")
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=20) as s:
            s.set_debuglevel(0)
            s.starttls(); s.login(SMTP_USER, SMTP_PASS); s.send_message(msg)
        print(f"Email successfully sent to {to_email}"); return True
    except smtplib.SMTPAuthenticationError as e: print(f"SMTP Auth Error for {SMTP_USER}: {e}\n{traceback.format_exc()}"); return False
    except smtplib.SMTPConnectError as e: print(f"SMTP Connect Error to {SMTP_SERVER}:{SMTP_PORT}: {e}\n{traceback.format_exc()}"); return False
    except smtplib.SMTPServerDisconnected as e: print(f"SMTP Server Disconnected: {e}\n{traceback.format_exc()}"); return False
    except smtplib.SMTPException as e: print(f"General SMTP Exception to {to_email}: {e}\n{traceback.format_exc()}"); return False
    except Exception as e: print(f"Unexpected error sending email to {to_email}: {e}\n{traceback.format_exc()}"); return False

# --- Syllabus & Lesson Plan Generation (Instructor Panel) ---
def generate_syllabus(cfg):
    sd, ed = datetime.strptime(cfg['start_date'], '%Y-%m-%d').date(), datetime.strptime(cfg['end_date'], '%Y-%m-%d').date()
    mr, total = f"{sd.strftime('%B')}–{ed.strftime('%B')}", count_classes(sd, ed, [days_map[d] for d in cfg['class_days']])
    header = [f"Course: {cfg['course_name']}", f"Prof: {cfg['instructor']['name']}", f"Email: {cfg['instructor']['email']}", f"Duration: {mr} ({total} classes)", '_'*60]
    objectives = [f" • {o}" for o in cfg['learning_objectives']]
    body = ["DESC:", cfg['course_description'], "", "OBJECTIVES:"] + objectives + ["", "GRADING:", " • Quiz per class.", " • Retake if <60%.", " • Final = quiz avg.", "", "SCHEDULE:", f" • {mr}, {', '.join(cfg['class_days'])}", "", "SUPPORT:", " • Office Hours: Tue 3–5PM; Thu 10–11AM (Zoom)", " • Email reply <24h weekdays"]
    return "\n".join(header + [""] + body)

def generate_plan_by_week_structured_and_formatted(cfg):
    sd, ed = datetime.strptime(cfg['start_date'], '%Y-%m-%d').date(), datetime.strptime(cfg['end_date'], '%Y-%m-%d').date()
    wdays = {days_map[d] for d in cfg['class_days']}
    class_dates = [cur for cur in (sd + timedelta(i) for i in range((ed - sd).days + 1)) if cur.weekday() in wdays]
    print(f"DEBUG: Class dates: {len(class_dates)}")
    if not class_dates: return "No class dates.", []
    full_text, char_map = cfg.get("full_text_content", ""), cfg.get("char_offset_page_map", [])
    if not full_text.strip():
        print("Warning: Full text empty.");
        placeholder_lessons, placeholder_lines, weeks_ph = [], [], {}
        for idx, dt in enumerate(class_dates):
            wk_key = f"{dt.isocalendar()[0]}-W{dt.isocalendar()[1]:02d}"
            ld = {"lesson_number": idx + 1, "date": dt.strftime('%Y-%m-%d'), "topic_summary": "Topic TBD (No PDF text)", "original_section_title": "N/A", "page_reference": None}
            placeholder_lessons.append(ld); weeks_ph.setdefault(wk_key, []).append(ld)
        for wk_key in sorted(weeks_ph.keys()):
            yr, wk = wk_key.split("-W"); placeholder_lines.append(f"**Week {wk} (Year {yr})**\n")
            for lsn in weeks_ph[wk_key]: placeholder_lines.append(f"**Lesson {lsn['lesson_number']} ({datetime.strptime(lsn['date'], '%Y-%m-%d').strftime('%B %d, %Y')})**: {lsn['topic_summary']}")
            placeholder_lines.append('')
        return "\n".join(placeholder_lines), placeholder_lessons

    total_chars, num_lessons = len(full_text), len(class_dates)
    chars_per_lesson = total_chars // num_lessons if num_lessons > 0 else total_chars
    min_chars, summaries, cur_ptr, seg_starts = 150, [], 0, []
    print(f"DEBUG: Total chars: {total_chars}, Chars/lesson: {chars_per_lesson}")
    for i in range(num_lessons):
        seg_starts.append(cur_ptr); start = cur_ptr
        end = cur_ptr + chars_per_lesson if i < num_lessons - 1 else total_chars
        seg_text, cur_ptr = full_text[start:end].strip(), end
        if len(seg_text) < min_chars: summaries.append("Review or brief topic.")
        else:
            try:
                print(f"DEBUG: Summarizing seg {i+1} (len {len(seg_text)}): '{seg_text[:70].replace(chr(10),' ')}...'")
                resp = openai.chat.completions.create(model="gpt-3.5-turbo", messages=[{"role":"system","content":"Identify core concept. Respond ONLY with short phrase (max 10-12 words) as lesson topic title, preferably gerund (e.g., 'Using verbs'). NO full sentences."}, {"role":"user","content": seg_text}], temperature=0.4, max_tokens=30)
                summaries.append(resp.choices[0].message.content.strip().replace('"', '').capitalize())
            except Exception as e: print(f"Error summarizing seg {i+1}: {e}"); summaries.append(f"Topic seg {i+1} (Summary Error)")

    lessons_by_course_week = {}
    structured_lessons = []
    if not class_dates: return "No class dates to process.", []
    course_week_counter = 0
    current_week_monday_for_grouping = None

    for idx, dt_obj in enumerate(class_dates):
        monday_of_this_week = dt_obj - timedelta(days=dt_obj.weekday())
        if current_week_monday_for_grouping is None or monday_of_this_week > current_week_monday_for_grouping:
            course_week_counter += 1
            current_week_monday_for_grouping = monday_of_this_week
        year_of_this_course_week = monday_of_this_week.year
        course_week_key = f"{year_of_this_course_week}-CW{course_week_counter:02d}"
        summary_for_lesson = summaries[idx]
        est_pg = None
        if char_map:
            seg_start = seg_starts[idx]
            for offset, pg in reversed(char_map):
                if seg_start >= offset: est_pg = pg; break
            if est_pg is None and char_map: est_pg = char_map[0][1]
        lesson_data = {
            "lesson_number": idx + 1,
            "date": dt_obj.strftime('%Y-%m-%d'),
            "topic_summary": summary_for_lesson,
            "original_section_title": f"Text Segment {idx+1}",
            "page_reference": est_pg
        }
        structured_lessons.append(lesson_data)
        lessons_by_course_week.setdefault(course_week_key, []).append(lesson_data)

    formatted_lines = []
    for course_week_key in sorted(lessons_by_course_week.keys()):
        year_disp, course_week_num_disp_str = course_week_key.split("-CW")
        course_week_num_disp = int(course_week_num_disp_str)
        first_date_in_this_week_group = lessons_by_course_week[course_week_key][0]['date']
        first_date_obj = datetime.strptime(first_date_in_this_week_group, '%Y-%m-%d')
        formatted_lines.append(f"**Course Week {course_week_num_disp} (Year {first_date_obj.year})**\n")
        for lesson in lessons_by_course_week[course_week_key]:
            ds = datetime.strptime(lesson['date'], '%Y-%m-%d').strftime('%B %d, %Y')
            pstr = f" (Approx. Ref. p. {lesson['page_reference']})" if lesson['page_reference'] else ''
            formatted_lines.append(f"**Lesson {lesson['lesson_number']} ({ds})**{pstr}: {lesson['topic_summary']}")
        formatted_lines.append('')
    return "\n".join(formatted_lines), structured_lessons

def generate_access_token(student_id, course_id, lesson_id, lesson_date_obj=None):
    access_code = generate_5_digit_code()  # Add this
    now = datetime.now(dt_timezone.utc)
    exp = now + timedelta(hours=LINK_VALIDITY_HOURS)
    payload = {
        "sub": student_id,
        "course_id": course_id,
        "lesson_id": lesson_id,
        "code": access_code,  # Include access code in the token payload
        "iat": now,
        "exp": exp,
        "aud": APP_DOMAIN
    }
    return jwt.encode(payload, JWT_SECRET_KEY, algorithm=ALGORITHM), access_code

def send_daily_class_reminders():
    print(f"SCHEDULER: Running daily class reminder job at {datetime.now(dt_timezone.utc)}")
    today_utc = datetime.now(dt_timezone.utc).date()
    for config_file in CONFIG_DIR.glob("*_config.json"):
        try:
            cfg = json.loads(config_file.read_text(encoding="utf-8"))
            course_id, course_name = config_file.stem.replace("_config", ""), cfg.get("course_name", "N/A")
            if not cfg.get("lessons") or not cfg.get("students"): continue
            for lesson in cfg["lessons"]:
                lesson_date = datetime.strptime(lesson["date"], '%Y-%m-%d').date()
                if lesson_date == today_utc:
                    print(f"SCHEDULER: Class found for {course_name} today: Lesson {lesson['lesson_number']}")
                    class_code = generate_5_digit_code() # This code isn't used in the current email template for link access
                    for student in cfg["students"]:
                        student_id, student_email, student_name = student.get("id", "unknown"), student.get("email"), student.get("name", "Student")
                        if not student_email: continue
                        token, access_code = generate_access_token(student_id, course_id, lesson["lesson_number"], lesson_date)
                        access_link = f"{APP_DOMAIN}/class?token={token}" # Use token[0] which is the JWT string
                        email_subject = f"Today's Class Link for {course_name}: {lesson['topic_summary']}"
                        email_html_body = f"""
                        <html><head><style>body {{font-family: sans-serif;}} strong {{color: #007bff;}} a {{color: #0056b3;}} .container {{padding: 20px; border: 1px solid #ddd; border-radius: 5px;}} .code {{font-size: 1.5em; font-weight: bold; background-color: #f0f0f0; padding: 5px 10px;}}</style></head>
                        <body><div class="container">
                            <p>Hi {student_name},</p>
                            <p>Your class for <strong>{course_name}</strong> - "{lesson['topic_summary']}" - is today!</p>
                            <p><strong>Your access code is:</strong> <span class="code">{access_code}</span></p>
                            <p>Access link: <a href="{access_link}">{access_link}</a></p>
                            <p>The link and code are valid for {LINK_VALIDITY_HOURS} hours from generation, typically covering morning to early afternoon UTC on {today_utc.strftime('%B %d, %Y')}.</p>
                            <p>Best regards,<br>AI Tutor System</p>
                        </div></body></html>"""
                        send_email_notification(student_email, email_subject, email_html_body, student_name) # from_name should be student_name if you want Reply-To to be student
        except Exception as e: print(f"SCHEDULER: Error in daily reminders for {config_file.name}: {e}\n{traceback.format_exc()}")

def log_student_progress(student_id, course_id, lesson_id, quiz_score_str, session_duration_secs, engagement_notes="N/A"):
    if not PROGRESS_LOG_FILE.exists():
        with open(PROGRESS_LOG_FILE, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["timestamp", "student_id", "course_id", "lesson_id", "quiz_score", "session_duration_seconds", "engagement_notes"])
    with open(PROGRESS_LOG_FILE, 'a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow([datetime.now(dt_timezone.utc).isoformat(), student_id, course_id, lesson_id, quiz_score_str, session_duration_secs, engagement_notes])
    print(f"Progress logged for student {student_id}, lesson {lesson_id} of course {course_id}.")

def check_student_progress_and_notify_professor():
    print(f"SCHEDULER: Running student progress check at {datetime.now(dt_timezone.utc)}")
    if not PROGRESS_LOG_FILE.exists():
        print("SCHEDULER: Progress log file does not exist. Skipping check.")
        return
    one_day_ago = datetime.now(dt_timezone.utc) - timedelta(days=1)
    alerts_to_send = {}
    try:
        with open(PROGRESS_LOG_FILE, 'r', newline='', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                try:
                    log_timestamp = datetime.fromisoformat(row["timestamp"])
                    if log_timestamp < one_day_ago: continue
                    quiz_score_str = row.get("quiz_score", "0/0")
                    if "/" in quiz_score_str:
                        correct, total_qs = map(int, quiz_score_str.split('/'))
                        if total_qs > 0 and (correct / total_qs) < 0.60:
                            student_id, course_id, lesson_id = row["student_id"], row["course_id"], row["lesson_id"]
                            alert_msg = (f"Student {student_id} scored {quiz_score_str} "
                                         f"on lesson {lesson_id} (logged {log_timestamp.strftime('%Y-%m-%d %H:%M')} UTC). "
                                         f"Session duration: {row.get('session_duration_seconds','N/A')}s. "
                                         f"Notes: {row.get('engagement_notes','N/A')}")
                            alerts_to_send.setdefault(course_id, {}).setdefault(student_id, []).append(alert_msg)
                except ValueError: print(f"SCHEDULER: Skipping malformed row in progress log: {row}"); continue
    except Exception as e_read_log: print(f"SCHEDULER: Error reading progress log: {e_read_log}"); return
    for course_id, student_alerts in alerts_to_send.items():
        config_path = CONFIG_DIR / f"{course_id}_config.json"
        if config_path.exists():
            try:
                cfg = json.loads(config_path.read_text(encoding="utf-8"))
                instructor_email = cfg.get("instructor", {}).get("email")
                instructor_name = cfg.get("instructor", {}).get("name", "Instructor")
                course_name = cfg.get("course_name", course_id)
                if instructor_email:
                    for student_id, messages in student_alerts.items():
                        subject = f"Student Progress Alert: {student_id} in {course_name}"
                        alerts_html = "<ul>" + "".join([f"<li>{msg}</li>" for msg in messages]) + "</ul>"
                        body_html = (f"<html><body><p>Dear {instructor_name},</p>"
                                     f"<p>One or more students in your course '{course_name}' may require attention based on recent AI Tutor sessions:</p>"
                                     f"{alerts_html}"
                                     f"<p>Please consider reviewing their progress and engaging with them directly.</p>"
                                     f"<p>Best regards,<br>AI Tutor Monitoring System</p></body></html>")
                        send_email_notification(instructor_email, subject, body_html, "AI Tutor System")
                        print(f"SCHEDULER: Sent progress alert for student {student_id} in course {course_name} to {instructor_email}")
            except Exception as e_send_alert: print(f"SCHEDULER: Error sending progress alert for course {course_id}: {e_send_alert}")

# --- Gradio Callbacks (Instructor Panel) ---
def _get_syllabus_text_from_config(course_name_str):
    if not course_name_str: return "Error: Course name missing."
    path = CONFIG_DIR / f"{course_name_str.replace(' ','_').lower()}_config.json"
    if not path.exists(): return f"Error: Config for '{course_name_str}' not found."
    try: return generate_syllabus(json.loads(path.read_text(encoding="utf-8")))
    except Exception as e: return f"Error loading syllabus: {e}"

def _get_plan_text_from_config(course_name_str):
    if not course_name_str: return "Error: Course name missing."
    path = CONFIG_DIR / f"{course_name_str.replace(' ','_').lower()}_config.json"
    if not path.exists(): return f"Error: Config for '{course_name_str}' not found."
    try: return json.loads(path.read_text(encoding="utf-8")).get("lesson_plan_formatted", "Plan not generated.")
    except Exception as e: return f"Error loading plan: {e}"

def enable_edit_syllabus_and_reload(current_course_name, current_output_content):
    if not current_output_content.strip().startswith("Course:"):
        syllabus_text = _get_syllabus_text_from_config(current_course_name)
        return gr.update(value=syllabus_text, interactive=True)
    return gr.update(interactive=True)

def enable_edit_plan_and_reload(current_course_name_for_plan, current_plan_output_content):
    if not current_plan_output_content.strip().startswith("**Week") and \
       (current_plan_output_content.strip().startswith("✅") or \
        current_plan_output_content.strip().startswith("⚠️")):
        plan_text = _get_plan_text_from_config(current_course_name_for_plan)
        return gr.update(value=plan_text, interactive=True)
    return gr.update(interactive=True)

def save_setup(course_name, instr_name, instr_email, devices, pdf_file, sy, sm, sd_day, ey, em, ed_day, class_days_selected, students_input_str):
    def error_return_tuple(error_message_str):
        return (gr.update(value=error_message_str, visible=True, interactive=False), gr.update(visible=True), None, gr.update(visible=False), gr.update(visible=False), gr.update(visible=False), gr.update(visible=False), gr.update(visible=False), gr.update(visible=False), gr.update(visible=False), gr.update(value="", visible=False), gr.update(visible=True), gr.update(visible=False))
    try:
        if not all([course_name, instr_name, instr_email, pdf_file, sy, sm, sd_day, ey, em, ed_day, class_days_selected]): return error_return_tuple("⚠️ Error: All fields marked with * are required.")
        try:
            start_dt, end_dt = datetime(int(sy), int(sm), int(sd_day)), datetime(int(ey), int(em), int(ed_day))
            if end_dt <= start_dt: return error_return_tuple("⚠️ Error: End date must be after start date.")
        except ValueError: return error_return_tuple("⚠️ Error: Invalid date selected.")

        sections_for_desc_obj = split_sections(pdf_file)
        if not sections_for_desc_obj or (len(sections_for_desc_obj) == 1 and "Error" in sections_for_desc_obj[0]['title']):
             return error_return_tuple("⚠️ Error: Could not extract structural sections from PDF for analysis.")

        full_pdf_text, char_offset_to_page_map, current_char_offset = "", [], 0
        fitz_available_for_full_text = fitz_available
        if fitz_available_for_full_text:
            doc_for_full_text = None
            try:
                if hasattr(pdf_file, "seek"): pdf_file.seek(0)
                if hasattr(pdf_file, "name"): doc_for_full_text = fitz.open(pdf_file.name)
                elif hasattr(pdf_file, "read"):
                    pdf_bytes = pdf_file.read(); pdf_file.seek(0)
                    doc_for_full_text = fitz.open(stream=pdf_bytes, filetype="pdf")
                if doc_for_full_text:
                    for page_num_fitz, page_obj in enumerate(doc_for_full_text):
                        page_text = page_obj.get_text("text", sort=True)
                        if page_text: char_offset_to_page_map.append((current_char_offset, page_num_fitz + 1)); full_pdf_text += page_text + "\n"; current_char_offset += len(page_text) + 1
                    doc_for_full_text.close()
                else: fitz_available_for_full_text = False
            except Exception as e_fitz_full: print(f"Error extracting full text with fitz: {e_fitz_full}"); fitz_available_for_full_text = False

        if not fitz_available_for_full_text or not full_pdf_text.strip():
            print("Warning: Fitz failed or not used for full text extraction, using concatenated sections. Page map will be empty or less accurate.")
            if hasattr(pdf_file, "seek"): pdf_file.seek(0)
            temp_sections = split_sections(pdf_file)
            full_pdf_text = "\n".join(s['content'] for s in temp_sections); char_offset_to_page_map = []

        if not full_pdf_text.strip(): return error_return_tuple("⚠️ Error: Extracted PDF text is empty.")

        full_content_for_ai_desc = "\n\n".join(f"Title: {s['title']}\nSnippet: {s['content'][:1000]}" for s in sections_for_desc_obj)
        r1 = openai.chat.completions.create(model="gpt-3.5-turbo", messages=[{"role":"system","content":"Generate a concise course description (2-3 sentences)."},{"role":"user","content": full_content_for_ai_desc}])
        desc = r1.choices[0].message.content.strip()
        r2 = openai.chat.completions.create(model="gpt-3.5-turbo", messages=[{"role":"system","content":"Generate 5–10 clear, actionable learning objectives. Start each with a verb."},{"role":"user","content": full_content_for_ai_desc}])
        objs = [ln.strip(" -•*") for ln in r2.choices[0].message.content.splitlines() if ln.strip()]
        parsed_students = [{"id": str(uuid.uuid4()), "name": n.strip(), "email": e.strip()} for ln in students_input_str.splitlines() if ',' in ln for n, e in [ln.split(',', 1)]]
        cfg = {"course_name": course_name, "instructor": {"name": instr_name, "email": instr_email}, "class_days": class_days_selected, "start_date": f"{sy}-{sm}-{sd_day}", "end_date": f"{ey}-{em}-{ed_day}", "allowed_devices": devices, "students": parsed_students, "sections_for_description": sections_for_desc_obj, "full_text_content": full_pdf_text, "char_offset_page_map": char_offset_to_page_map, "course_description": desc, "learning_objectives": objs, "lessons": [], "lesson_plan_formatted": ""}
        path = CONFIG_DIR / f"{course_name.replace(' ','_').lower()}_config.json"
        path.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
        syllabus_text = generate_syllabus(cfg)
        return (gr.update(value=syllabus_text, visible=True, interactive=False), gr.update(visible=False), None, gr.update(visible=True), gr.update(visible=True), gr.update(visible=True), gr.update(visible=False), gr.update(visible=False), gr.update(visible=True), gr.update(visible=True), gr.update(value="", visible=False), gr.update(visible=False), gr.update(visible=True, value=course_name))
    except openai.APIError as oai_err: print(f"OpenAI Error: {oai_err}\n{traceback.format_exc()}"); return error_return_tuple(f"⚠️ OpenAI API Error: {oai_err}.")
    except Exception as e: print(f"Error in save_setup: {e}\n{traceback.format_exc()}"); return error_return_tuple(f"⚠️ Error: {e}")

def generate_plan_callback(course_name_from_input):
    def error_return_for_plan(error_message_str):
        return (
            gr.update(value=error_message_str, visible=True, interactive=False),
            None, None,
            gr.update(visible=True),
            None, None,
            gr.update(visible=False),
            gr.update(visible=False),
        )

    try:
        if not course_name_from_input:
            return error_return_for_plan("⚠️ Error: Course Name required.")

        config_path = CONFIG_DIR / f"{course_name_from_input.replace(' ','_').lower()}_config.json"
        if not config_path.exists():
            return error_return_for_plan(f"⚠️ Error: Config for '{course_name_from_input}' not found.")

        # Load config and generate lesson plan
        cfg = json.loads(config_path.read_text(encoding="utf-8"))
        formatted_plan_str, structured_lessons_list = generate_plan_by_week_structured_and_formatted(cfg)
        cfg["lessons"] = structured_lessons_list
        cfg["lesson_plan_formatted"] = formatted_plan_str
        config_path.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")

        # ——— FIRST-DAY EMAIL LOGIC ———
        today_iso    = date.today().isoformat()
        first_lesson = cfg["lessons"][0] if cfg["lessons"] else None
        print(f"DEBUG [generate_plan]: today={today_iso}, lesson1 date={first_lesson['date'] if first_lesson else None}")

        if first_lesson and first_lesson["date"] == today_iso:
            for student_info in cfg["students"]:
                token, access_code = generate_access_token(
                    student_info["id"],
                    course_name_from_input.replace(" ", "_").lower(),
                    first_lesson["lesson_number"],
                    datetime.strptime(first_lesson["date"], "%Y-%m-%d").date()
                )
                access_link = f"{APP_DOMAIN}/class?token={token}"

                print(f"DEBUG [generate_plan]: sending email to {student_info['email']} → {access_link}")

                html_body = f"""
                <html><head><style>body {{font-family: sans-serif;}} strong {{color: #007bff;}} a {{color: #0056b3;}} .container {{padding: 20px; border: 1px solid #ddd; border-radius: 5px;}} .code {{font-size: 1.5em; font-weight: bold; background-color: #f0f0f0; padding: 5px 10px;}}</style></head>
                <body><div class="container">
                    <p>Hi {student_info['name']},</p>
                    <p>Your course <strong>{cfg['course_name']}</strong> starts <strong>today</strong>!</p>
                    <p><strong>Your access code is:</strong> <span class="code">{access_code}</span></p>
                    <p>Access link: <a href="{access_link}">{access_link}</a></p>
                    <p>The link and code are valid for {LINK_VALIDITY_HOURS} hours from generation.</p>
                    <p>Good luck!<br>AI Tutor System</p>
                </div></body></html>
                """

                sent = send_email_notification(
                    student_info["email"],
                    f"{cfg['course_name']} — Your Class Link for Today",
                    html_body
                )
                print(f"DEBUG [generate_plan]: email sent to {student_info['email']}? {sent}")
        # ————————————————————————————————

        # Build UI notification
        class_days_str = ", ".join(cfg.get("class_days", ["configured schedule"]))
        notification_message = (
            f"\n\n---\n"
            f"✅ **Lesson Plan Generated & Email System Activated for Class Days!**\n"
            f"Students will now receive a unique link on each scheduled day ({class_days_str}), "
            f"valid {LINK_VALIDITY_HOURS} hours from generation."
        )
        display_text_for_plan_box = formatted_plan_str + notification_message

        return (
            gr.update(value=display_text_for_plan_box, visible=True, interactive=False),
            None, None,
            gr.update(visible=False),
            None, None,
            gr.update(visible=True),
            gr.update(visible=True),
        )

    except openai.APIError as oai_err:
        print(f"OpenAI Error: {oai_err}\n{traceback.format_exc()}")
        return error_return_for_plan(f"⚠️ OpenAI API Error: {oai_err}.")
    except Exception as e:
        print(f"Error in generate_plan_callback: {e}\n{traceback.format_exc()}")
        return error_return_for_plan(f"⚠️ Error: {e}")


        # Build the “lesson plan generated” message for the UI
        class_days_str = ", ".join(cfg.get("class_days", ["configured schedule"]))
        notification_message = (
            f"\n\n---\n"
            f"✅ **Lesson Plan Generated & Email System Activated for Class Days!**\n"
            f"Students will now receive a unique link on each scheduled day ({class_days_str}), "
            f"valid {LINK_VALIDITY_HOURS} hours from generation."
        )
        display_text_for_plan_box = formatted_plan_str + notification_message

        return (
            gr.update(value=display_text_for_plan_box, visible=True, interactive=False),
            None, None,
            gr.update(visible=False),
            None, None,
            gr.update(visible=True),
            gr.update(visible=True),
        )

    except openai.APIError as oai_err:
        print(f"OpenAI Error: {oai_err}\n{traceback.format_exc()}")
        return error_return_for_plan(f"⚠️ OpenAI API Error: {oai_err}.")
    except Exception as e:
        print(f"Error in generate_plan_callback: {e}\n{traceback.format_exc()}")
        return error_return_for_plan(f"⚠️ Error: {e}")


def email_document_callback(course_name, doc_type, output_text_content, students_input_str):
    if not SMTP_USER or not SMTP_PASS: return gr.update(value="⚠️ Error: SMTP settings not configured.")
    try:
        if not course_name or not output_text_content: return gr.update(value=f"⚠️ Error: Course Name & {doc_type} content required.")
        path = CONFIG_DIR / f"{course_name.replace(' ','_').lower()}_config.json"
        if not path.exists(): return gr.update(value=f"⚠️ Error: Config for '{course_name}' not found.")
        cfg = json.loads(path.read_text(encoding="utf-8")); instr_name, instr_email = cfg.get("instructor", {}).get("name", "Instructor"), cfg.get("instructor", {}).get("email")
        
        # Create a temporary file-like object for the attachment
        buf, fn = download_docx(output_text_content, f"{course_name.replace(' ','_')}_{doc_type.lower()}.docx")
        # For send_email_notification, we need to pass an object that has a .name attribute
        # and can be opened in 'rb' mode if it's a path.
        # Or, pass the BytesIO buffer directly if send_email_notification is adapted.
        # Let's assume send_email_notification can handle a BytesIO object if it has a .name
        # If not, we'd save buf to a temp file and pass that path.
        # For simplicity with the current send_email_notification, we'll simulate a file object with .name
        class TempFileLike:
            def __init__(self, buffer, name):
                self.buffer = buffer
                self.name = name # This is the filename for the attachment
                self.buffer.seek(0)
            def read(self): # Mimic read method if needed by mimetypes or email lib
                return self.buffer.read()
            def seek(self, offset):
                self.buffer.seek(offset)

        # Create a named temporary file to pass to send_email_notification
        temp_file_path = Path(fn) # Use a temporary path
        with open(temp_file_path, 'wb') as tmp_f:
            tmp_f.write(buf.getvalue())
        
        # Create a file object mock if needed for send_email_notification's attachment logic
        class MockFile:
            def __init__(self, path):
                self.name = str(path) # Path to the actual temp file

        attachment_obj_for_email = MockFile(temp_file_path)


        recipients = ([{"name":instr_name, "email":instr_email}] if instr_email else []) + [{"name":n.strip(), "email":e.strip()} for ln in students_input_str.splitlines() if ',' in ln for n,e in [ln.split(',',1)]]
        if not recipients: return gr.update(value="⚠️ Error: No recipients.")
        
        s_count = 0
        errs = []
        
        # Prepare HTML content once
        html_email_body = f"""
        <html><body>
        <p>Hi {{recipient_name}},</p>
        <p>Attached is the {doc_type.lower()} for the course: <strong>{course_name}</strong>.</p>
        <p>Best regards,<br>AI Tutor System</p>
        </body></html>
        """

        for rec in recipients:
            personalized_html_body = html_email_body.replace("{{recipient_name}}", rec['name'])
            subject = f"{doc_type.capitalize()}: {course_name}"
            
            # Pass the MockFile object which has the .name attribute pointing to the temp file
            if send_email_notification(rec["email"], subject, personalized_html_body, from_name=SMTP_USER, attachment_file_obj=attachment_obj_for_email):
                s_count += 1
            else:
                errs.append(f"Failed to send to {rec['email']}. Check logs for SMTP errors.")

        if temp_file_path.exists(): # Clean up the temporary file
            os.remove(temp_file_path)

        status = f"✅ {doc_type.capitalize()} sent attempt to {s_count} recipient(s)."
        if errs: status += f"\n⚠️ Errors:\n" + "\n".join(errs)
        return gr.update(value=status)

    except Exception as e: err_txt = f"⚠️ Emailing Err:\n{traceback.format_exc()}"; print(err_txt); return gr.update(value=err_txt)


def email_syllabus_callback(c, s_str, out_content): return email_document_callback(c, "Syllabus", out_content, s_str)
def email_plan_callback(c, s_str, out_content): return email_document_callback(c, "Lesson Plan", out_content, s_str)

# --- Build Instructor UI ---
def build_instructor_ui():
    import time
    with gr.Blocks(theme=gr.themes.Soft()) as instructor_demo:
        gr.Markdown("## AI Tutor Instructor Panel")
        with gr.Tabs():
            with gr.TabItem("Course Setup & Syllabus"):
                with gr.Row():
                    course = gr.Textbox(label="Course Name*")
                    instr = gr.Textbox(label="Instructor Name*")
                    email = gr.Textbox(label="Instructor Email*", type="email")
                pdf_file = gr.File(label="Upload Course Material PDF*", file_types=[".pdf"])
                with gr.Row():
                    with gr.Column(scale=2):
                        gr.Markdown("#### Course Schedule")
                        years = [str(y) for y in range(datetime.now().year, datetime.now().year + 5)]
                        months = [f"{m:02d}" for m in range(1, 13)]
                        days_list = [f"{d:02d}" for d in range(1, 32)]
                        with gr.Row(): sy, sm, sd_day = gr.Dropdown(years, label="Start Year*"), gr.Dropdown(months, label="Start Month*"), gr.Dropdown(days_list, label="Start Day*")
                        with gr.Row(): ey, em, ed_day = gr.Dropdown(years, label="End Year*"), gr.Dropdown(months, label="End Month*"), gr.Dropdown(days_list, label="End Day*")
                        class_days_selected = gr.CheckboxGroup(list(days_map.keys()), label="Class Days*")
                    with gr.Column(scale=1):
                        gr.Markdown("#### Student & Access")
                        devices = gr.CheckboxGroup(["Phone", "PC", "Tablet"], label="Allowed Devices", value=["PC"])
                        students_input_str = gr.Textbox(label="Students (Name,Email per line)", lines=5, placeholder="S. One,s1@ex.com\nS. Two,s2@ex.com")
                btn_save = gr.Button("1. Save Setup & Generate Syllabus", variant="primary")
                gr.Markdown("---")
                output_box = gr.Textbox(label="Output", lines=20, interactive=False, visible=False, show_copy_button=True)
                with gr.Row(visible=False) as syllabus_actions_row:
                    btn_edit_syl = gr.Button(value="📝 Edit Syllabus Text")
                    btn_email_syl = gr.Button(value="📧 Email Syllabus", variant="secondary")
            with gr.TabItem("Lesson Plan Management"):
                lesson_plan_setup_message = gr.Markdown(value="### Course Setup Required\nCourse Setup (on Tab 1) must be completed before generating a Lesson Plan.", visible=True)
                course_load_for_plan = gr.Textbox(label="Course Name for Lesson Plan", placeholder="e.g., Introduction to Python", visible=False)
                output_plan_box = gr.Textbox(label="Lesson Plan Output", lines=20, interactive=False, visible=False, show_copy_button=True)
                with gr.Row(visible=False) as plan_buttons_row:
                    btn_generate_plan = gr.Button("2. Generate/Re-generate Lesson Plan", variant="primary")
                    btn_edit_plan = gr.Button(value="📝 Edit Plan Text")
                    btn_email_plan = gr.Button(value="📧 Email Lesson Plan", variant="secondary")
            with gr.TabItem("Contact Support"):
                gr.Markdown("### Send a Message to Support")
                with gr.Row(): contact_name, contact_email_addr = gr.Textbox(label="Your Name"), gr.Textbox(label="Your Email Address")
                contact_message = gr.Textbox(label="Message", lines=5, placeholder="Type your message here...")
                contact_attachment = gr.File(label="Attach File (Optional)", file_count="single")
                btn_send_contact_email = gr.Button("Send Message", variant="primary")
                contact_status_output = gr.Markdown(value="")
                def handle_contact_submission(name, email_addr, message_text, attachment_file):
                    errors = []
                    if not name.strip(): errors.append("Name is required.")
                    if not email_addr.strip(): errors.append("Email is required.")
                    elif "@" not in email_addr: errors.append("Enter a valid email.")
                    if not message_text.strip(): errors.append("Message is required.")
                    if errors: return (gr.update(value="Please fix:\n" + "\n".join(f"- {e}" for e in errors)), gr.update(value=name), gr.update(value=email_addr), gr.update(value=message_text), gr.update(value=attachment_file))
                    sent = send_email_notification(SUPPORT_EMAIL_ADDRESS, f"Contact: {name} <{email_addr}>", message_text.replace("\n","<br>"), from_name=email_addr, attachment_file_obj=attachment_file)
                    if sent: return (gr.update(value="<span style='color:green;'>Sent! ✔</span>"), gr.update(value=""), gr.update(value=""), gr.update(value=""), gr.update(value=None))
                    else: return (gr.update(value="<span style='color:red;'>Failed to send. Check SMTP or logs.</span>"), gr.update(value=name), gr.update(value=email_addr), gr.update(value=message_text), gr.update(value=attachment_file))
                btn_send_contact_email.click(handle_contact_submission, inputs=[contact_name, contact_email_addr, contact_message, contact_attachment], outputs=[contact_status_output, contact_name, contact_email_addr, contact_message, contact_attachment], queue=True)
        dummy_btn_1, dummy_btn_2, dummy_btn_3, dummy_btn_4 = gr.Button(visible=False), gr.Button(visible=False), gr.Button(visible=False), gr.Button(visible=False)
        btn_save.click(save_setup, inputs=[course, instr, email, devices, pdf_file, sy, sm, sd_day, ey, em, ed_day, class_days_selected, students_input_str], outputs=[output_box, btn_save, dummy_btn_1, btn_generate_plan, btn_edit_syl, btn_email_syl, btn_edit_plan, btn_email_plan, syllabus_actions_row, plan_buttons_row, output_plan_box, lesson_plan_setup_message, course_load_for_plan])
        btn_edit_syl.click(enable_edit_syllabus_and_reload, inputs=[course, output_box], outputs=[output_box])
        btn_email_syl.click(email_syllabus_callback, inputs=[course, students_input_str, output_box], outputs=[output_box])
        btn_generate_plan.click(generate_plan_callback, inputs=[course_load_for_plan], outputs=[output_plan_box, dummy_btn_2, dummy_btn_1, btn_generate_plan, dummy_btn_3, dummy_btn_4, btn_edit_plan, btn_email_plan]).then(lambda: (gr.update(visible=True), gr.update(visible=True)), outputs=[output_plan_box, plan_buttons_row])
        btn_edit_plan.click(enable_edit_plan_and_reload, inputs=[course_load_for_plan, output_plan_box], outputs=[output_plan_box])
        btn_email_plan.click(email_plan_callback, inputs=[course_load_for_plan, students_input_str, output_plan_box], outputs=[output_plan_box])
        course.change(lambda x: x, inputs=[course], outputs=[course_load_for_plan])
    return instructor_demo

# --- Student System Prompt Generation (Placeholder) ---
def generate_student_system_prompt(mode, interests, topic, segment_text):
    # This is a placeholder. Implement your actual prompt generation logic.
    # The more detailed and context-aware this prompt is, the better the tutor will behave.
    prompt = f"You are {STUDENT_BOT_NAME}, a friendly and encouraging AI Tutor.\n"
    prompt += f"Your current interaction mode is: '{mode}'.\n"
    prompt += f"The student's known interests are: {interests if interests else 'not yet specified'}.\n"
    prompt += f"Today's lesson topic is: '{topic}'.\n"
    prompt += f"Focus on this text segment for context (if relevant to mode): '{segment_text[:700]}...' (truncated for brevity in this example prompt)\n\n"

    if mode == "initial_greeting":
        prompt += "Greet the student warmly. Briefly introduce yourself and the lesson topic. Ask an open-ended question to start the conversation, perhaps about their hobbies or what they already know about the topic. Keep your first message concise."
    elif mode == "onboarding":
        prompt += "Continue the onboarding process. Engage with the student's response about their interests. Ask another question to learn more or smoothly transition towards the lesson content. Show genuine curiosity."
    elif mode == "teaching_transition":
        prompt += "Smoothly transition from the general chat/onboarding to the core lesson content. You could say something like, 'That's fascinating! Speaking of [relate to topic if possible], let's dive into our main topic for today: {topic}.'"
    elif mode == "teaching":
        prompt += f"Focus on teaching the lesson based on '{topic}' and the provided text segment. Explain concepts clearly, provide examples, and ask checking questions to ensure understanding. Relate to student interests if possible."
    elif mode == "interest_break_transition":
        prompt += "You've been teaching for a bit. Announce a short 'interest break'. You could say, 'Alright, let's take a quick breather from the main lesson. You mentioned you like [student_interest]. Tell me more about that!' or ask a fun, light question related to their interests."
    elif mode == "interest_break_active":
        prompt += f"Engage in a short, light-hearted conversation based on the student's interests: {interests}. After one or two turns, gently guide the conversation back to the lesson topic."
    elif mode == "quiz_time":
        prompt += f"It's time for a quick quiz question on what you've learned about '{topic}'. Ask one clear multiple-choice or short-answer question based on the material. After the student answers, provide feedback (correct/incorrect and a brief explanation)."
    elif mode == "ending_session":
        prompt += "The session is ending. Briefly summarize what was covered. Thank the student for their participation. If there was a quiz, mention their performance positively. Wish them well."
    else: # Default fallback
        prompt += "Continue the conversation naturally, keeping the lesson goals in mind. Be supportive and engaging."

    prompt += "\nYour responses should be tailored to a student with an English level of B1 (Intermediate). Use clear language. Be encouraging and patient."
    return prompt

# --- Student Tutor UI and Logic ---
def build_student_tutor_ui():
    with gr.Blocks(theme=gr.themes.Soft()) as student_demo:
        # --- Persistent states ---
        token_state          = gr.State(None)
        course_id_state      = gr.State(None)
        lesson_id_state      = gr.State(None)
        student_id_state     = gr.State(None)
        lesson_topic_state   = gr.State(None)
        lesson_segment_state = gr.State(None)
        st_chat_history      = gr.State([])
        st_display_history   = gr.State([])
        st_student_profile   = gr.State({"interests": [], "quiz_score": {"correct": 0, "total": 0}, "english_level": STUDENT_DEFAULT_ENGLISH_LEVEL})
        st_session_mode      = gr.State("initial_greeting")
        st_turn_count        = gr.State(0)
        st_teaching_turns    = gr.State(0)
        st_session_start     = gr.State(None)

        # --- UI Header ---
        gr.Markdown("## Easy AI Tutor - Interactive Lesson")

        # --- Student Input and Output Interface ---
        with gr.Row():
            with gr.Column(scale=1):
                st_voice_dropdown = gr.Dropdown(choices=["alloy", "echo", "fable", "nova", "onyx", "shimmer"], value="nova", label="Tutor Voice")
                st_mic_input = gr.Audio(sources=["microphone"], type="filepath", label="🎤 Record your answer")
                st_text_input = gr.Textbox(label="💬 Or type your response", placeholder="Type here...")
                st_send_button = gr.Button("Send", variant="primary")
            with gr.Column(scale=3):
                st_chatbot = gr.Chatbot(label="Lesson Conversation", height=500, bubble_full_width=False)
                st_audio_out = gr.Audio(type="filepath", autoplay=True, label="🎧 Tutor’s Voice")

        # --- Extract token from query ---
        def grab_token(request: gr.Request):
            return request.query_params.get("token")

        student_demo.load(fn=grab_token, inputs=[], outputs=[token_state])

        # --- Decode token and load context ---
        def decode_context(token, request: gr.Request):
            """
            • Декодирует JWT-токен и проверяет 5-значный код из URL
            • Загружает конфиг курса
            • Возвращает: course_id, lesson_id, student_id, topic, segment_title
              ─ topic никогда не пустой и всегда приведён к Title Case
            """
            print(f"DEBUG: decode_context called. Token present: {bool(token)}, Code from URL: {request.query_params.get('code')}")
        
            # 1) Токен должен быть
            if not token:
                print("DEBUG: decode_context returning - No Token")
                return (
                    "Unknown Course", "N/A", "Unknown Student",
                    "Error: No Token",
                    "Please ensure you accessed this page via a valid link."
                )
        
            try:
                # 2) Проверяем подпись токена и код
                payload = jwt.decode(
                    token,
                    JWT_SECRET_KEY,
                    algorithms=[ALGORITHM],
                    audience=APP_DOMAIN
                )
                print(f"DEBUG: JWT Payload decoded: {payload}")
                if payload.get("code") != request.query_params.get("code"):
                    print(f"DEBUG: decode_context returning - Code Mismatch. Expected {payload.get('code')}, got {request.query_params.get('code')}")
                    return (
                        "N/A", "N/A", "N/A",
                        "Error: Code Mismatch",
                        "Access code mismatch. Please recheck the link or code."
                    )
        
                # 3) Базовые ID
                course_id  = payload["course_id"]
                student_id = payload["sub"]
                # Ensure lesson_id is an int, handle potential errors
                try:
                    lesson_id  = int(payload["lesson_id"])
                except (ValueError, TypeError):
                    print(f"DEBUG: decode_context returning - Lesson ID Invalid Format in payload. lesson_id: {payload.get('lesson_id')}")
                    return (
                        course_id, "N/A", student_id,
                        "Error: Lesson ID Invalid Format",
                        "Lesson ID in token is not a valid number."
                    )
                print(f"DEBUG: Extracted from payload - course_id: {course_id}, student_id: {student_id}, lesson_id: {lesson_id} (type: {type(lesson_id)})")
        
                # 4) Загружаем конфиг курса
                cfg_path = CONFIG_DIR / f"{course_id}_config.json"
                print(f"DEBUG: Attempting to load config from: {cfg_path}")
                if not cfg_path.exists():
                    print(f"DEBUG: decode_context returning - Course Config Missing. Path: {cfg_path}")
                    return (
                        course_id, lesson_id, student_id,
                        "Error: Course Config Missing",
                        f"No config file found for this course ({course_id})."
                    )
        
                cfg       = json.loads(cfg_path.read_text(encoding="utf-8"))
                lessons   = cfg.get("lessons", [])
                print(f"DEBUG: Config loaded. Number of lessons found: {len(lessons)}")
        
                if not isinstance(lesson_id, int) or lesson_id <= 0 or lesson_id > len(lessons):
                    print(f"DEBUG: decode_context returning - Lesson Invalid. lesson_id: {lesson_id}, num_lessons: {len(lessons)}")
                    return (
                        course_id, lesson_id if isinstance(lesson_id, int) else "N/A", student_id,
                        "Error: Lesson Invalid",
                        f"Lesson ID ({lesson_id}) is out of range or invalid for {len(lessons)} lessons."
                    )
        
                # 5) Извлекаем тему и сегмент
                lesson_index = lesson_id - 1
                lesson = lessons[lesson_index]
                print(f"DEBUG: Accessing lesson at index {lesson_index}. Lesson data: {lesson}")
        
                topic_summary_raw = lesson.get("topic_summary")
                topic_raw = lesson.get("topic")
                title_raw = lesson.get("title")
                name_raw = lesson.get("name")
                print(f"DEBUG: Raw topic fields - topic_summary: '{topic_summary_raw}', topic: '{topic_raw}', title: '{title_raw}', name: '{name_raw}'")
        
                current_topic = (
                    topic_summary_raw
                    or topic_raw
                    or title_raw
                    or name_raw
                    # Fallback if all primary fields are empty or None
                    or (f"Lesson {lesson_id}" if (topic_summary_raw is None and topic_raw is None and title_raw is None and name_raw is None) else None)
                )
                
                if not current_topic or not current_topic.strip(): # If current_topic is None or empty string after OR chain
                     current_topic = f"Lesson {lesson_id} (Default Topic)"
                     print(f"DEBUG: All specific topic fields were empty/None. Using default: {current_topic}")
        
        
                # ► Нормализуем к Title Case («Learning Spanish Greetings»)
                # Ensure current_topic is a string before calling .title()
                current_topic_title_cased = str(current_topic).title() if current_topic else f"Lesson {lesson_id} (Topic Processing Error)"
                print(f"DEBUG: Final current_topic before return: '{current_topic_title_cased}'")
        
                current_segment = lesson.get("segment_title") or lesson.get("original_section_title") or ""
                print(f"DEBUG: Final current_segment before return: '{current_segment}'")
        
                # 6) Возвращаем 5 значений
                print(f"DEBUG: decode_context successfully returning: {(course_id, lesson_id, student_id, current_topic_title_cased, current_segment)}")
                return (
                    course_id,
                    lesson_id,
                    student_id,
                    current_topic_title_cased,
                    current_segment
                )
        
            # 7) Обработка ошибок токена
            except jwt.ExpiredSignatureError:
                print("DEBUG: decode_context returning - ExpiredSignatureError")
                return (
                    "N/A", "N/A", "N/A",
                    "Error: Expired",
                    "This link has expired."
                )
            except jwt.InvalidTokenError as e:
                print(f"DEBUG: decode_context returning - InvalidTokenError: {e}")
                return (
                    "N/A", "N/A", "N/A",
                    "Error: Invalid Token",
                    f"Invalid token: {e}"
                )
            except Exception as e:
                # Try to get some context if payload was partially decoded
                course_id_fallback = "N/A"
                student_id_fallback = "N/A"
                lesson_id_fallback = "N/A"
                if 'payload' in locals() and isinstance(payload, dict):
                    course_id_fallback = payload.get("course_id", "N/A")
                    student_id_fallback = payload.get("sub", "N/A")
                    lesson_id_fallback = payload.get("lesson_id", "N/A")
        
                print(f"DEBUG: decode_context returning - Unknown Exception: {e}, Traceback: {traceback.format_exc()}")
                return (
                    course_id_fallback, lesson_id_fallback, student_id_fallback,
                    "Error: Unknown Processing",
                    f"Unexpected error during context decoding: {e}"
                )

        decode_event = student_demo.load(
            fn=decode_context,
            inputs=[token_state],
            outputs=[
                course_id_state,
                lesson_id_state,
                student_id_state,
                lesson_topic_state,
                lesson_segment_state,
            ],
        )


        # --- Initial tutor message ---
        # In build_student_tutor_ui() function:

        # --- Initial tutor message ---
        # Add lesson_id to function arguments
        def tutor_greeter(current_lesson_topic, current_lesson_segment, current_lesson_id,
                  request: gr.Request): # Add request
            # A. Handle cases where the topic itself is an error message from decode_context
            if isinstance(current_lesson_topic, str) and current_lesson_topic.startswith("Error:"):
                error_message_for_ui = f"⚠️ **Access Problem:** {current_lesson_topic.replace('Error: ', '')}.\n"
                if "Expired" in current_lesson_topic:
                    error_message_for_ui += "Your session link may have expired. Please try obtaining a new link."
                elif "Token" in current_lesson_topic or "Code Mismatch" in current_lesson_topic:
                    error_message_for_ui += "Please ensure you are using the correct and complete link, including any access code."
                else:
                    error_message_for_ui += "Please contact support or your instructor if this issue persists."
        
                # Return values to display the error and disable interaction
                return (
                    [[None, error_message_for_ui]],  # st_display_history (for st_chatbot)
                    [],                              # st_chat_history (empty, no AI interaction)
                    "error",                         # st_session_mode
                    0,                               # st_turn_count
                    0,                               # st_teaching_turns
                    None,                            # st_audio_out (no audio for error)
                    datetime.now(dt_timezone.utc),   # st_session_start
                    gr.update(interactive=False),    # st_mic_input (disable)
                    gr.update(interactive=False),    # st_text_input (disable)
                    gr.update(interactive=False)     # st_send_button (disable)
                )
        
            # B. Proceed with normal greeting if topic is valid
            lesson_id_str = str(current_lesson_id) if current_lesson_id is not None else "?"
            # Ensure display_topic is not None or empty, fallback if necessary
            display_topic = current_lesson_topic if current_lesson_topic and current_lesson_topic.strip() else f"Lesson {lesson_id_str} (Topic Undefined)"
        
            # Ensure segment is a string, default to empty if not
            current_lesson_segment = current_lesson_segment if isinstance(current_lesson_segment, str) else ""
        
            prompt = generate_student_system_prompt("initial_greeting", "", display_topic, current_lesson_segment)
            
            audio_fp_str = None
            msg_content = ""
        
            try:
                client = openai.OpenAI()
                res = client.chat.completions.create(
                    model=STUDENT_CHAT_MODEL,
                    messages=[{"role": "system", "content": prompt}],
                    max_tokens=150
                )
                msg_content = res.choices[0].message.content.strip()
                
                try:
                    # Use the selected voice from st_voice_dropdown if available, otherwise default
                    # For the initial greeting, we don't have access to st_voice_dropdown's value directly here
                    # So we'll stick to a default like "nova" or make it configurable if needed.
                    # For simplicity, keeping "nova" for now.
                    speech_res = client.audio.speech.create(model=STUDENT_TTS_MODEL, voice="nova", input=msg_content)
                    audio_fp = STUDENT_AUDIO_DIR / f"intro_{uuid.uuid4()}.mp3"
                    with open(audio_fp, "wb") as f:
                        f.write(speech_res.content)
                    audio_fp_str = str(audio_fp)
                except Exception as e_tts:
                    print(f"TTS Error in tutor_greeter for main response: {e_tts}")
        
            except Exception as e_chat:
                print(f"Chat Completion Error in tutor_greeter: {e_chat}")
                msg_content = f"Hello! I'm ready to start our lesson on '{display_topic}', but I'm having a slight technical difficulty with my opening lines. How are you today?"
                try:
                    client_fallback_tts = openai.OpenAI() 
                    speech_res_fallback = client_fallback_tts.audio.speech.create(model=STUDENT_TTS_MODEL, voice="nova", input=msg_content)
                    audio_fp_fallback = STUDENT_AUDIO_DIR / f"intro_fallback_{uuid.uuid4()}.mp3"
                    with open(audio_fp_fallback, "wb") as f:
                        f.write(speech_res_fallback.content)
                    audio_fp_str = str(audio_fp_fallback)
                except Exception as e_tts_fallback:
                    print(f"TTS Error in tutor_greeter for fallback message: {e_tts_fallback}")
        
            initial_display_history = [[None, msg_content]]
            initial_chat_history = [{"role": "system", "content": prompt}, {"role": "assistant", "content": msg_content}]
            
            # Return values for a successful greeting, inputs remain interactive
            return (
                initial_display_history,
                initial_chat_history,
                "onboarding", # st_session_mode
                0,            # st_turn_count
                0,            # st_teaching_turns
                audio_fp_str,
                datetime.now(dt_timezone.utc),
                gr.update(interactive=True), # st_mic_input
                gr.update(interactive=True), # st_text_input
                gr.update(interactive=True)  # st_send_button
            )

        # Update the student_demo.load call for tutor_greeter to include lesson_id_state
        decode_event.then(
            tutor_greeter,
            inputs=[
                lesson_topic_state,
                lesson_segment_state,
                lesson_id_state
                # gr.Request() REMOVED FROM HERE
            ],
            outputs=[
                st_display_history, # This will be used by st_chatbot
                st_chat_history,
                st_session_mode,
                st_turn_count,
                st_teaching_turns,
                st_audio_out,
                st_session_start,
                st_mic_input,     # Add output for mic
                st_text_input,    # Add output for text input
                st_send_button    # Add output for send button
            ],
        )
        # --- Processing student response ---
        def handle_response(mic_path, text, chat_hist, disp_hist, profile, mode, turns, teaching_turns, voice,
                            sid, cid, lid, topic, segment, start_time):
            input_text = text.strip() if text else ""
            if mic_path:
                try:
                    client = openai.OpenAI()
                    with open(mic_path, "rb") as f:
                        result = client.audio.transcriptions.create(file=f, model=STUDENT_WHISPER_MODEL)
                    input_text = result.text.strip()
                    if os.path.exists(mic_path): os.remove(mic_path)
                except:
                    input_text = "(Audio could not be transcribed.)"

            if not input_text:
                return disp_hist, chat_hist, profile, mode, turns, teaching_turns, None, gr.update(value=None), gr.update(value="")

            disp_hist.append([input_text, None])
            chat_hist.append({"role": "user", "content": input_text})

            turns += 1
            if mode == "onboarding":
                profile["interests"].append(input_text)
                if turns >= STUDENT_ONBOARDING_TURNS:
                    mode = "teaching_transition"
            elif mode == "teaching_transition":
                mode = "teaching"
            elif mode == "teaching":
                teaching_turns += 1
                if teaching_turns % STUDENT_TEACHING_TURNS_PER_BREAK == 0:
                    mode = "interest_break_transition"
            elif mode == "interest_break_transition":
                mode = "interest_break_active"
            elif mode == "interest_break_active":
                mode = "teaching"
            elif turns >= STUDENT_MAX_SESSION_TURNS:
                mode = "ending_session"

            prompt = generate_student_system_prompt(mode, ", ".join(profile["interests"]), topic, segment)
            if chat_hist and chat_hist[0]["role"] != "system":
                chat_hist.insert(0, {"role": "system", "content": prompt})

            try:
                client = openai.OpenAI()
                res = client.chat.completions.create(model=STUDENT_CHAT_MODEL, messages=chat_hist, max_tokens=250)
                bot_reply = res.choices[0].message.content.strip()
            except:
                bot_reply = "Sorry, I didn't understand that. Could you rephrase?"

            chat_hist.append({"role": "assistant", "content": bot_reply})
            disp_hist[-1][1] = bot_reply

            try:
                speech = client.audio.speech.create(model=STUDENT_TTS_MODEL, voice=voice, input=bot_reply)
                fp = STUDENT_AUDIO_DIR / f"turn_{uuid.uuid4()}.mp3"
                with open(fp, "wb") as f:
                    f.write(speech.content)
                return disp_hist, chat_hist, profile, mode, turns, teaching_turns, str(fp), gr.update(value=None), gr.update(value="")
            except:
                return disp_hist, chat_hist, profile, mode, turns, teaching_turns, None, gr.update(value=None), gr.update(value="")

        event_inputs = [
            st_mic_input, st_text_input, st_chat_history, st_display_history, st_student_profile,
            st_session_mode, st_turn_count, st_teaching_turns, st_voice_dropdown,
            student_id_state, course_id_state, lesson_id_state, lesson_topic_state, lesson_segment_state, st_session_start
        ]
        event_outputs = [
            st_chatbot, st_chat_history, st_student_profile, st_session_mode,
            st_turn_count, st_teaching_turns, st_audio_out, st_mic_input, st_text_input
        ]

        st_mic_input.change(fn=handle_response, inputs=event_inputs, outputs=event_outputs)
        st_text_input.submit(fn=handle_response, inputs=event_inputs, outputs=event_outputs)
        st_send_button.click(fn=handle_response, inputs=event_inputs, outputs=event_outputs)

    return student_demo

# --- FastAPI App Setup (Continued) ---

# Mount your Gradio Instructor UI under /instructor
instructor_ui = build_instructor_ui()
app = gr.mount_gradio_app(app, instructor_ui, path="/instructor")

# Build and Mount Student Tutor UI
student_tutor_ui_instance = build_student_tutor_ui()
app = gr.mount_gradio_app(app, student_tutor_ui_instance, path=STUDENT_UI_PATH)


# Redirect root (/) → /instructor so users just type your domain
# Verification step before showing student tutor interface
@app.get("/verify_access", response_class=HTMLResponse)
async def verify_access(request: Request, token: str = None):
    if not token:
        return HTMLResponse("<h3>Error: Missing token. Please use your lesson link.</h3>", status_code=400)
    return HTMLResponse(f"""
    <html><head><title>Access Verification</title></head>
    <body style="font-family:sans-serif; margin:50px;">
        <h2>Enter Your Access Code</h2>
        <form method="get" action="/class/enter">
            <input type="hidden" name="token" value="{token}">
            <input type="text" name="code" placeholder="5-digit code" required>
            <button type="submit">Continue</button>
        </form>


    </body>
    </html>
    """)

@app.get("/")
def root_redirect(): # Renamed to avoid conflict if you define root differently elsewhere
    return RedirectResponse(url="/instructor")

# Endpoint for student to access lesson via token
@app.get("/class", response_class=HTMLResponse)
async def get_student_lesson_page(request: Request, token: str = None): # token is the JWT string
    if not token:
        return HTMLResponse("<h3>Error: Access token missing. Please use the link provided in your email.</h3>", status_code=400)
    try:
        # Validate the token's basic structure, signature, and expiry. Audience is also checked.
        # This decode is primarily for validation before redirecting to manual code entry.
        # The full payload processing will happen in the student UI's decode_context after code entry.
        jwt.decode(token, JWT_SECRET_KEY, algorithms=[ALGORITHM], audience=APP_DOMAIN)

        # Redirect to the verification page, passing the original token.
        # The /verify_access page will ask the student to input the 5-digit code from their email.
        verify_url = f"/verify_access?token={token}"
        return RedirectResponse(url=verify_url)

    except jwt.ExpiredSignatureError:
        return HTMLResponse("<h3>Access link has expired.</h3><p>Your session link was valid for a limited time. Please check if a new link is available or contact your instructor.</p>", status_code=401)
    except jwt.InvalidTokenError as e:
        # This covers various issues like InvalidSignatureError, InvalidAudienceError, DecodeError etc.
        print(f"Invalid token error on /class: {e}")
        return HTMLResponse(f"<h3>Invalid access link.</h3><p>There was a problem with your session link: {str(e)}. Please ensure you copied the entire link correctly.</p>", status_code=401)
    except Exception as e:
        print(f"Error processing /class request: {e}\n{traceback.format_exc()}")
        return HTMLResponse(f"<h3>Error preparing lesson.</h3><p>An unexpected error occurred: {str(e)}. Please try again later or contact support.</p>", status_code=500)

# ─── Final gate: token + code must match ────────────────────────────────────
@app.get("/class/enter", response_class=RedirectResponse)
async def enter_class(token: str, code: str):
    """
    Last check before we let the student see the Gradio UI:
    1. JWT must be valid (signature, expiry, audience)
    2. 5-digit code in the URL must match the one stored in the token
    """
    try:
        payload = jwt.decode(
            token,
            JWT_SECRET_KEY,
            algorithms=[ALGORITHM],
            audience=APP_DOMAIN
        )

        if payload.get("code") != code:
            return HTMLResponse("<h3>Wrong access code.</h3>", status_code=401)

        # Everything checks out → hand off to the real UI
        return RedirectResponse(
            url=f"{STUDENT_UI_PATH}?token={token}&code={code}"
        )

    except jwt.ExpiredSignatureError:
        return HTMLResponse("<h3>Link expired.</h3>", status_code=401)
    except jwt.InvalidTokenError:
        return HTMLResponse("<h3>Invalid link.</h3>", status_code=401)
    except Exception as e:
        return HTMLResponse(f"<h3>Unexpected error: {e}</h3>", status_code=500)

@app.on_event("startup")
async def startup_event():
    scheduler.add_job(send_daily_class_reminders, trigger=CronTrigger(hour=5, minute=50, timezone='UTC'), id="daily_reminders", name="Daily Class Reminders", replace_existing=True)
    scheduler.add_job(check_student_progress_and_notify_professor, trigger=CronTrigger(hour=18, minute=0, timezone='UTC'), id="progress_check", name="Student Progress Check", replace_existing=True)
    if not scheduler.running:
        try:
            scheduler.start()
            print("APScheduler started successfully.")
        except Exception as e:
            print(f"Error starting APScheduler: {e}")
    else:
        print("APScheduler already running.")
    for job in scheduler.get_jobs(): print(f"  Job: {job.id}, Name: {job.name}, Trigger: {job.trigger}")

@app.on_event("shutdown")
async def shutdown_event():
    if scheduler.running:
        scheduler.shutdown()
        print("APScheduler shutdown.")

if __name__ == "__main__":
    print(f"Starting App. Instructor Panel at /instructor. Student access via /class?token=...")
    print(f"Student Tutor UI should be available at {STUDENT_UI_PATH} (after /class redirect)")
    # To run this: uvicorn your_script_name:app --reload --port 8000
    # Example: uvicorn main:app --reload --port 8000
    # (Assuming your file is named main.py)
    #
    # The uvicorn command handles running the FastAPI app and its lifecycle.
    # No need for a while True loop here.
    pass
